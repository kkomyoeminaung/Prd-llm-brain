"""
Document parsers for various file formats
Supports: PDF, DOCX, TXT, HTML, ZIP
"""

import os
import io
import zipfile
import tempfile
from typing import List, Dict, Optional, Any
from dataclasses import dataclass
import re


@dataclass
class DocumentChunk:
    """A chunk of document with metadata"""
    content: str
    source: str
    page: int = 0
    chunk_index: int = 0
    metadata: Dict = None


class TextParser:
    """Parse plain text files"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    @staticmethod
    def parse_bytes(content: bytes) -> str:
        return content.decode('utf-8', errors='ignore')


class PDFParser:
    """Parse PDF files"""
    
    @staticmethod
    def parse(file_path: str) -> List[Dict]:
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    pages.append({'page': i + 1, 'content': text})
            return pages
        except ImportError:
            print("[PDFParser] Install pypdf: !pip install pypdf")
            return []
    
    @staticmethod
    def parse_bytes(content: bytes) -> List[Dict]:
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=True) as tmp:
            tmp.write(content)
            tmp.flush()
            return PDFParser.parse(tmp.name)


class DOCXParser:
    """Parse DOCX files"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            return '\n'.join(full_text)
        except ImportError:
            print("[DOCXParser] Install python-docx: !pip install python-docx")
            return ""
    
    @staticmethod
    def parse_bytes(content: bytes) -> str:
        import docx
        doc = docx.Document(io.BytesIO(content))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        return '\n'.join(full_text)


class HTMLParser:
    """Parse HTML files"""
    
    @staticmethod
    def parse(file_path: str) -> str:
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                return text
        except ImportError:
            print("[HTMLParser] Install beautifulsoup4: !pip install beautifulsoup4")
            return ""
    
    @staticmethod
    def parse_bytes(content: bytes) -> str:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        return '\n'.join(line for line in lines if line)


class URLParser:
    """Fetch and parse content from URLs"""
    
    @staticmethod
    async def fetch(url: str) -> Optional[str]:
        import aiohttp
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=aiohttp.ClientTimeout(total=30)) as resp:
                    if resp.status == 200:
                        html = await resp.text()
                        
                        # Parse HTML
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(html, 'html.parser')
                        for script in soup(["script", "style", "nav", "footer", "header"]):
                            script.decompose()
                        text = soup.get_text()
                        lines = (line.strip() for line in text.splitlines())
                        return '\n'.join(line for line in lines if line)
        except Exception as e:
            print(f"[URLParser] Error fetching {url}: {e}")
        return None


class ZIPParser:
    """Extract and parse ZIP files"""
    
    @staticmethod
    def parse(file_path: str) -> List[Dict]:
        results = []
        with zipfile.ZipFile(file_path, 'r') as zf:
            for file_info in zf.filelist:
                if file_info.filename.endswith(('.txt', '.pdf', '.docx', '.html')):
                    with zf.open(file_info) as f:
                        content = f.read()
                        parsed = DocumentIngestor.parse_bytes(content, file_info.filename)
                        if parsed:
                            results.append({
                                'filename': file_info.filename,
                                'content': parsed
                            })
        return results
    
    @staticmethod
    def parse_bytes(content: bytes) -> List[Dict]:
        with tempfile.NamedTemporaryFile(suffix='.zip', delete=True) as tmp:
            tmp.write(content)
            tmp.flush()
            return ZIPParser.parse(tmp.name)


class DocumentIngestor:
    """Main document ingestion orchestrator"""
    
    @staticmethod
    def parse_file(file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt':
            return TextParser.parse(file_path)
        elif ext == '.pdf':
            pages = PDFParser.parse(file_path)
            return '\n'.join(p['content'] for p in pages)
        elif ext == '.docx':
            return DOCXParser.parse(file_path)
        elif ext == '.html':
            return HTMLParser.parse(file_path)
        elif ext == '.zip':
            files = ZIPParser.parse(file_path)
            return '\n'.join(f['content'] for f in files)
        else:
            return f"Unsupported file type: {ext}"
    
    @staticmethod
    def parse_bytes(content: bytes, filename: str) -> str:
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == '.txt':
            return TextParser.parse_bytes(content)
        elif ext == '.pdf':
            pages = PDFParser.parse_bytes(content)
            return '\n'.join(p['content'] for p in pages)
        elif ext == '.docx':
            return DOCXParser.parse_bytes(content)
        elif ext == '.html':
            return HTMLParser.parse_bytes(content)
        elif ext == '.zip':
            files = ZIPParser.parse_bytes(content)
            return '\n'.join(f['content'] for f in files)
        else:
            return f"Unsupported file type: {ext}"
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk:
                chunks.append(chunk)
        
        return chunks
